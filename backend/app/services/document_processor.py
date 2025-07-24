import io
import logging
from typing import Dict, Any, Optional
import PyPDF2
import pdfplumber
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("PyMuPDF not available, using fallback PDF readers")
from docx import Document as DocxDocument
from app.models.document import DocumentType
from app.services.ocr_service import ocr_service

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing uploaded documents (PDF, Word)."""
    
    async def extract_text(self, file_content: bytes, file_type: str) -> str:
        """Extract text from document based on file type."""
        
        if file_type.lower() == 'pdf':
            return await self._extract_pdf_text(file_content)
        elif file_type.lower() in ['docx', 'doc']:
            return await self._extract_docx_text(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        
        text = ""
        logger.info(f"Starting PDF extraction. File size: {len(file_content)} bytes")
        
        # Try PyMuPDF first if available (best for Japanese text)
        if PYMUPDF_AVAILABLE:
            try:
                pdf_document = fitz.open(stream=file_content, filetype="pdf")
                logger.info(f"PyMuPDF: Successfully opened PDF with {pdf_document.page_count} pages")
                
                for page_num in range(pdf_document.page_count):
                    try:
                        page = pdf_document[page_num]
                        # Try multiple text extraction methods
                        page_text = page.get_text()
                        
                        # Also try getting text with different options
                        if not page_text or not page_text.strip():
                            page_text = page.get_text("text")  # explicit text format
                        
                        if page_text and page_text.strip():
                            # Ensure proper encoding
                            if isinstance(page_text, bytes):
                                page_text = page_text.decode('utf-8', errors='replace')
                            text += page_text + "\n"
                            logger.info(f"PyMuPDF Page {page_num+1}: extracted {len(page_text)} characters")
                        else:
                            # Check if page has any content at all
                            page_dict = page.get_text("dict")
                            blocks = page_dict.get("blocks", [])
                            image_count = len(page.get_images())
                            
                            # Try to get text from annotations (form fields)
                            annot_text = ""
                            for annot in page.annots():
                                if annot.info["content"]:
                                    annot_text += annot.info["content"] + " "
                            
                            if annot_text.strip():
                                text += annot_text + "\n"
                                logger.info(f"PyMuPDF Page {page_num+1}: extracted {len(annot_text)} characters from annotations")
                            else:
                                logger.warning(f"PyMuPDF Page {page_num+1}: no text extracted. Blocks: {len(blocks)}, Images: {image_count}")
                    except Exception as page_error:
                        logger.warning(f"PyMuPDF Page {page_num+1} extraction failed: {str(page_error)}")
                
                pdf_document.close()
                
                if text.strip():
                    logger.info(f"PyMuPDF SUCCESS: Total extracted {len(text)} characters")
                    return text.strip()
                else:
                    logger.warning("PyMuPDF extracted no text, trying pdfplumber")
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {str(e)}, trying pdfplumber")
        
        # Try pdfplumber second (good for Japanese text)
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                logger.info(f"pdfplumber: Successfully opened PDF with {len(pdf.pages)} pages")
                for i, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            # Ensure proper encoding for Japanese text
                            if isinstance(page_text, bytes):
                                page_text = page_text.decode('utf-8', errors='replace')
                            text += page_text + "\n"
                            logger.info(f"pdfplumber Page {i+1}: extracted {len(page_text)} characters")
                        else:
                            logger.warning(f"pdfplumber Page {i+1}: no text extracted")
                    except Exception as page_error:
                        logger.warning(f"pdfplumber Page {i+1} extraction failed: {str(page_error)}")
            
            if text.strip():
                logger.info(f"pdfplumber SUCCESS: Total extracted {len(text)} characters")
                return text.strip()
            else:
                logger.warning("pdfplumber extracted no text, trying PyPDF2")
        except Exception as e:
            logger.warning(f"pdfplumber failed completely: {str(e)}, trying PyPDF2")
        
        # Fallback to PyPDF2
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
            logger.info(f"PyPDF2: Successfully opened PDF with {num_pages} pages")
            
            for page_num in range(num_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        # Ensure proper encoding
                        if isinstance(page_text, bytes):
                            page_text = page_text.decode('utf-8', errors='replace')
                        text += page_text + "\n"
                        logger.info(f"PyPDF2 Page {page_num+1}: extracted {len(page_text)} characters")
                    else:
                        logger.warning(f"PyPDF2 Page {page_num+1}: no text extracted")
                except Exception as page_error:
                    logger.warning(f"PyPDF2 Page {page_num+1} extraction failed: {str(page_error)}")
            
            if text.strip():
                logger.info(f"PyPDF2 SUCCESS: Total extracted {len(text)} characters")
                return text.strip()
            else:
                # If still no text, check if PDF might be scanned/image-based
                logger.warning("No text extracted from PDF by any method - attempting OCR")
                
                # Try OCR as last resort
                ocr_text = await ocr_service.extract_text_from_image_pdf(file_content)
                
                if ocr_text and ocr_text.strip():
                    logger.info(f"OCR SUCCESS: Extracted {len(ocr_text)} characters")
                    return ocr_text.strip()
                else:
                    logger.error("OCR also failed to extract text from PDF")
                    raise ValueError("PDFからテキストを抽出できませんでした。スキャンされたPDFの可能性があり、OCRも失敗しました。")
        except PyPDF2.errors.PdfReadError as e:
            logger.error(f"PDF is corrupted or encrypted: {str(e)}")
            raise ValueError("PDFファイルが破損しているか、暗号化されています。")
        except Exception as e:
            logger.error(f"PDF extraction failed completely: {str(e)}")
            raise ValueError(f"PDFの処理に失敗しました: {str(e)}")
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def detect_document_type(self, text: str, filename: str) -> DocumentType:
        """Detect document type based on content and filename."""
        
        text_lower = text.lower()
        filename_lower = filename.lower()
        
        # Check for resume indicators
        resume_indicators = ['履歴書', '氏名', '生年月日', '住所', '学歴', '職歴']
        if any(indicator in text for indicator in resume_indicators) or '履歴書' in filename_lower:
            return DocumentType.RESUME
        
        # Check for CV indicators
        cv_indicators = ['職務経歴書', '職務経歴', '業務内容', 'プロジェクト', '実績', 'スキル']
        if any(indicator in text for indicator in cv_indicators) or '職務経歴' in filename_lower:
            return DocumentType.CV
        
        # Check for skill sheet indicators
        skill_indicators = ['スキルシート', 'skill sheet', '技術スタック', '開発経験']
        if any(indicator in text_lower for indicator in skill_indicators):
            return DocumentType.SKILL_SHEET
        
        return DocumentType.OTHER
    
    def parse_japanese_resume(self, text: str) -> Dict[str, Any]:
        """Parse Japanese resume (履歴書) structure."""
        
        parsed = {
            "personal_info": {},
            "education": [],
            "work_history": [],
            "qualifications": [],
            "other_info": {}
        }
        
        # This is a simplified parser - in production, you'd use more sophisticated NLP
        lines = text.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect sections
            if '学歴' in line:
                current_section = 'education'
            elif '職歴' in line:
                current_section = 'work_history'
            elif '資格' in line or '免許' in line:
                current_section = 'qualifications'
            elif current_section:
                # Add to appropriate section
                if current_section == 'education':
                    parsed['education'].append(line)
                elif current_section == 'work_history':
                    parsed['work_history'].append(line)
                elif current_section == 'qualifications':
                    parsed['qualifications'].append(line)
        
        return parsed
    
    def parse_japanese_cv(self, text: str) -> Dict[str, Any]:
        """Parse Japanese CV (職務経歴書) structure."""
        
        parsed = {
            "summary": "",
            "work_experience": [],
            "projects": [],
            "skills": [],
            "achievements": []
        }
        
        # This is a simplified parser - in production, you'd use more sophisticated NLP
        # You might want to integrate with Gemini API for better parsing
        
        return parsed


# Singleton instance
document_processor = DocumentProcessor()